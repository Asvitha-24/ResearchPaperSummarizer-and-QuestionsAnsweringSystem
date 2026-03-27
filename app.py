#!/usr/bin/env python
"""
Flask API Server for Research Paper Summarizer & QA System
Connects the backend ML models with the frontend React application
"""

from flask import Flask, request, jsonify
from flask_cors import CORS
import os
import sys
from pathlib import Path
import threading
import time
import pickle
import hashlib

# Add src to path
sys.path.insert(0, os.path.dirname(__file__))

# Lazy import for ResearchPaperQASystem (to avoid slow transformers loading on startup)
ResearchPaperQASystem = None

# Try to import PDF libraries
try:
    from pypdf import PdfReader
    HAS_PDF = True
    print("[OK] pypdf library loaded successfully")
except ImportError:
    try:
        import PyPDF2
        HAS_PDF = True
        print("[OK] PyPDF2 library loaded")
    except ImportError:
        HAS_PDF = False
        print("[WARN] PDF extraction libraries not installed.")

# Try to import DOCX library
try:
    from docx import Document
    HAS_DOCX = True
    print("[OK] python-docx library loaded successfully")
except ImportError:
    HAS_DOCX = False
    print("[WARN] python-docx not installed. Install with: pip install python-docx")

# Model caching configuration
MODEL_CACHE_DIR = Path('checkpoints/model_cache')
MODEL_CACHE_DIR.mkdir(parents=True, exist_ok=True)
QA_SYSTEM_CACHE_FILE = MODEL_CACHE_DIR / 'qa_system_cache.pkl'
QA_SYSTEM_LOCK = threading.Lock()

# Initialize the QA System (eager loading to improve response times)
qa_system = None
qa_system_loading = False
qa_system_ready = False

def _get_model_version_hash() -> str:
    """Generate a hash of model configuration for versioning."""
    config = "bart-large-cnn_distilbert-squad_v1"
    return hashlib.md5(config.encode()).hexdigest()

def _save_qa_system_cache(qa_sys):
    """Save QA system to disk cache."""
    try:
        version_file = MODEL_CACHE_DIR / 'version.txt'
        with open(version_file, 'w') as f:
            f.write(_get_model_version_hash())
        
        # Note: Pickling large transformer models is not recommended due to size
        # Instead, we cache the model files themselves via HuggingFace cache
        print(f"[CACHE] Model cache directories created at {MODEL_CACHE_DIR}")
        return True
    except Exception as e:
        print(f"[WARN] Failed to save cache metadata: {e}")
        return False

def _check_model_cache_valid() -> bool:
    """Check if model cache is still valid."""
    try:
        version_file = MODEL_CACHE_DIR / 'version.txt'
        if not version_file.exists():
            return False
        
        with open(version_file, 'r') as f:
            cached_version = f.read().strip()
        
        return cached_version == _get_model_version_hash()
    except:
        return False

def initialize_qa_system_sync():
    """Initialize QA system synchronously (blocking)."""
    global qa_system, qa_system_ready
    
    print("\n" + "="*80)
    print("[SERVER STARTUP] Initializing ML models...")
    print("This may take 1-3 minutes on first start...")
    print("="*80)
    
    try:
        from src.model import ResearchPaperQASystem
        start_time = time.time()
        qa_system = ResearchPaperQASystem()
        elapsed = time.time() - start_time
        
        # Cache model configuration
        _save_qa_system_cache(qa_system)
        qa_system_ready = True
        
        print(f"\n✅ QA System initialized successfully in {elapsed:.1f}s")
        print("="*80 + "\n")
        return True
    except Exception as e:
        print(f"\n⚠️  Warning: QA System initialization failed: {e}")
        print("The system will load models on first use instead")
        print("="*80 + "\n")
        import traceback
        traceback.print_exc()
        qa_system = None
        qa_system_ready = False
        return False

def initialize_qa_system_background():
    """Initialize QA system in background thread (non-blocking)."""
    global qa_system, qa_system_loading, qa_system_ready
    
    def _load_in_background():
        global qa_system
        print("\n[BACKGROUND] Starting QA System initialization in background thread...")
        try:
            from src.model import ResearchPaperQASystem
            start_time = time.time()
            qa_system = ResearchPaperQASystem()
            elapsed = time.time() - start_time
            
            _save_qa_system_cache(qa_system)
            
            print(f"\n[BACKGROUND] ✅ QA System ready in {elapsed:.1f}s")
            global qa_system_ready
            qa_system_ready = True
        except Exception as e:
            print(f"\n[BACKGROUND] ⚠️  QA System initialization failed: {e}")
            qa_system = None
            qa_system_ready = False
    
    if not qa_system_loading:
        qa_system_loading = True
        thread = threading.Thread(target=_load_in_background, daemon=True)
        thread.start()

def get_qa_system():
    """Get the QA system (with fallback lazy loading if needed)."""
    global qa_system, qa_system_ready
    
    # If already loaded and ready, return it
    if qa_system is not None and qa_system_ready:
        return qa_system
    
    # If QA system is None and we haven't tried loading yet, try now
    if qa_system is None:
        with QA_SYSTEM_LOCK:
            # Double-check after acquiring lock
            if qa_system is None:
                try:
                    from src.model import ResearchPaperQASystem
                    print("[QA_SYSTEM] Lazy loading QA system on first request...")
                    start_time = time.time()
                    qa_system = ResearchPaperQASystem()
                    elapsed = time.time() - start_time
                    qa_system_ready = True
                    print(f"[QA_SYSTEM] ✅ Loaded in {elapsed:.1f}s")
                except Exception as e:
                    print(f"[QA_SYSTEM] ⚠️  Lazy loading failed: {e}")
                    qa_system = None
                    qa_system_ready = False
    
    return qa_system

# Initialize Flask app
app = Flask(__name__)
CORS(app)  # Enable CORS for frontend requests

# Initialize QA system at startup
# Options: 'sync' (blocks startup), 'background' (non-blocking), 'lazy' (on first use)
LOAD_STRATEGY = 'background'  # Default to background loading

print(f"\n[STARTUP] Using '{LOAD_STRATEGY}' model loading strategy")

if LOAD_STRATEGY == 'sync':
    print("Waiting for QA System initialization...")
    initialize_qa_system_sync()
elif LOAD_STRATEGY == 'background':
    print("Starting QA System initialization in background...")
    initialize_qa_system_background()
else:  # 'lazy'
    print("QA System will load on first request (lazy loading)")

# Helper function to extract text from PDF
def extract_pdf_text(filepath):
    """
    Extract text from PDF file using pypdf with aggressive word boundary detection
    """
    try:
        if not HAS_PDF:
            file_size = os.path.getsize(filepath)
            return f"[PDF File: {os.path.basename(filepath)}]\n[File size: {file_size / 1024:.2f} KB]\n[Note: PDF extraction not available. Please install: pip install pypdf]"
        
        text = []
        from pypdf import PdfReader
        
        with open(filepath, 'rb') as file:
            pdf_reader = PdfReader(file)
            
            # Extract text from all pages
            for page_num in range(len(pdf_reader.pages)):
                try:
                    page = pdf_reader.pages[page_num]
                    page_text = page.extract_text()
                    if page_text and page_text.strip():
                        text.append(page_text)
                except Exception as e:
                    print(f"Error extracting page {page_num}: {e}")
                    continue
        
        if not text:
            return f"[PDF file uploaded but no extractable text found in {os.path.basename(filepath)}]"
        
        pdf_content = '\n'.join(text)
        
        # Import regex module
        import re
        
        # AGGRESSIVE WORD BOUNDARY INSERTION
        # This handles concatenated text like "Thesystemwillgeneratehigh-qualitysummaries"
        
        # Step 1: Add spaces between lowercase and uppercase (CamelCase)
        # "Thesystem" -> "The system"
        pdf_content = re.sub(r'([a-z])([A-Z])', r'\1 \2', pdf_content)
        
        # Step 2: Add spaces before common English word starts
        common_words = [
            'the', 'and', 'for', 'with', 'that', 'from', 'have', 'this', 'will', 'your', 
            'about', 'which', 'their', 'would', 'been', 'has', 'were', 'does', 'how', 'what', 
            'when', 'where', 'why', 'all', 'can', 'did', 'get', 'let', 'may', 'new', 'now',
            'old', 'one', 'our', 'out', 'own', 'say', 'she', 'too', 'use', 'way', 'who',
            'boy', 'car', 'day', 'dog', 'get', 'go', 'good', 'hand', 'he', 'her', 'him',
            'is', 'it', 'just', 'know', 'like', 'look', 'make', 'man', 'many', 'me', 'my',
            'no', 'not', 'of', 'on', 'only', 'or', 'people', 'see', 'some', 'such', 'take',
            'than', 'them', 'then', 'there', 'these', 'they', 'thing', 'think', 'through',
            'to', 'tell', 'time', 'two', 'up', 'us', 'very', 'want', 'we', 'work', 'year',
            'you', 'after', 'also', 'as', 'at', 'be', 'because', 'been', 'before', 'by',
            'come', 'could', 'down', 'each', 'even', 'find', 'first', 'follow', 'give',
            'go', 'group', 'has', 'high', 'his', 'home', 'if', 'in', 'into', 'is', 'its',
            'just', 'keep', 'kind', 'last', 'life', 'little', 'live', 'long', 'love',
            'much', 'must', 'name', 'never', 'next', 'number', 'off', 'often', 'other',
            'over', 'part', 'place', 'play', 'present', 'problem', 'provide', 'put',
            'quality', 'question', 'quite', 'rather', 'read', 'result', 'right', 'run',
            'same', 'school', 'second', 'set', 'should', 'show', 'small', 'so', 'something',
            'sometimes', 'state', 'still', 'study', 'system', 'take', 'tell', 'than',
            'that', 'the', 'their', 'them', 'these', 'they', 'thing', 'think', 'this',
            'those', 'though', 'through', 'thus', 'to', 'today', 'together', 'too', 'took',
            'try', 'type', 'under', 'understand', 'unit', 'until', 'upon', 'use', 'used',
            'usually', 'value', 'various', 'very', 'view', 'want', 'was', 'water', 'way',
            'week', 'well', 'went', 'were', 'what', 'when', 'where', 'which', 'while',
            'who', 'whole', 'why', 'will', 'with', 'within', 'without', 'word', 'work',
            'world', 'would', 'write', 'written', 'year', 'yes', 'yet', 'you', 'young',
            'your', 'yourself'
        ]
        
        # Create pattern for common words at word boundary
        for word in common_words:
            # Match: (lowercase or digit)(common word in lowercase)
            # Case-insensitive to catch accidental lowercase
            pattern = rf'([a-z\d])({word})(?=[a-z])'
            pdf_content = re.sub(pattern, rf'\1 \2', pdf_content, flags=re.IGNORECASE)
        
        # Step 3: Handle digits followed by letters (like "5Proposed" -> "5 Proposed")
        pdf_content = re.sub(r'(\d)([a-zA-Z])', r'\1 \2', pdf_content)
        
        # Step 4: Handle letters followed by digits
        pdf_content = re.sub(r'([a-zA-Z])(\d)', r'\1 \2', pdf_content)
        
        # Step 5: Fix internal word spacing (co mp ress -> compress)
        pdf_content = re.sub(r'(\w)\s{1,2}(?=[a-z])', r'\1', pdf_content)
        
        # Step 6: Fix excessive whitespace
        pdf_content = re.sub(r'\n\s*\n', '\n\n', pdf_content)  # Keep paragraph breaks
        pdf_content = re.sub(r' {2,}', ' ', pdf_content)  # Multiple spaces to single
        pdf_content = re.sub(r'\t', ' ', pdf_content)  # Tabs to spaces
        
        # Step 7: Clean line breaks
        pdf_content = re.sub(r'-\n', '', pdf_content)  # Hyphenated breaks
        pdf_content = re.sub(r'(\w)\n(?=[a-z])', r'\1 ', pdf_content)  # Join line breaks mid-word
        pdf_content = re.sub(r'\n(?=[a-z])', ' ', pdf_content)  # Space after line breaks before lowercase
        
        # Step 8: Fix punctuation spacing
        pdf_content = re.sub(r'\s+([.,!?;:])', r'\1', pdf_content)  # No space before punctuation
        pdf_content = re.sub(r'([.!?])\s{1,}([A-Z])', r'\1 \2', pdf_content)  # Single space after
        
        # Step 9: Additional aggressive patterns for remaining concatenation
        # Handle academic terms and common research words that might stay concatenated
        research_terms = [
            'research', 'model', 'network', 'algorithm', 'learning', 'neural', 'data', 'analysis',
            'method', 'approach', 'framework', 'system', 'performance', 'result', 'evaluation',
            'training', 'testing', 'validation', 'accuracy', 'metric', 'parameter', 'optimization',
            'implementation', 'architecture', 'layer', 'feature', 'extraction', 'classification',
            'prediction', 'inference', 'inference', 'abstract', 'conclusion', 'introduction', 'method',
            'experiment', 'baseline', 'benchmark', 'dataset', 'corpus', 'annotation', 'evaluation'
        ]
        
        for term in research_terms:
            # Catch patterns like "proposedmodel" -> "proposed model"
            pattern = rf'([a-z])({term})(?=[a-z])'
            pdf_content = re.sub(pattern, rf'\1 \2', pdf_content, flags=re.IGNORECASE)
        
        # Step 10: Additional pass for all-lowercase concatenated words (very aggressive)
        # This catches "thesystemwill" -> "the system will" using word length heuristics
        import re
        words_to_check = pdf_content.split()
        processed_words = []
        
        for word in words_to_check:
            if len(word) > 15 and word.islower() and not any(c.isdigit() for c in word):
                # Very long lowercase word, likely concatenated - try to split it
                # Use a simple heuristic: common syllable patterns
                split_word = re.sub(r'(ing|tion|ness|ment|able|ible|ful|less|ly|er|or|ar|ian)(?=[a-z])', r'\1 ', word)
                if split_word != word:
                    word = split_word
            processed_words.append(word)
        
        pdf_content = ' '.join(processed_words)
        
        # Step 11: Final cleanup
        pdf_content = re.sub(r' {2,}', ' ', pdf_content)  # Multiple spaces to single
        pdf_content = '\n'.join(line.strip() for line in pdf_content.split('\n') if line.strip())
        
        return pdf_content if pdf_content else f"[No readable content in {os.path.basename(filepath)}]"
    
    except Exception as e:
        print(f"PDF extraction error: {e}")
        return f"[Error extracting PDF: {str(e)}]"

# Helper function to extract text from text files
def extract_text_file(filepath):
    """Extract text from plain text file"""
    try:
        with open(filepath, 'r', encoding='utf-8') as file:
            return file.read()
    except Exception as e:
        print(f"Text extraction error: {e}")
        return None

# Helper function to extract text from DOCX files
def extract_docx_text(filepath):
    """
    Extract text from DOCX file using python-docx with word boundary restoration
    """
    try:
        if not HAS_DOCX:
            file_size = os.path.getsize(filepath)
            return f"[DOCX File: {os.path.basename(filepath)}]\n[File size: {file_size / 1024:.2f} KB]\n[Note: DOCX extraction not available. Please install: pip install python-docx]"
        
        from docx import Document
        import re
        
        doc = Document(filepath)
        text_paragraphs = []
        
        # Extract text from all paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_paragraphs.append(paragraph.text)
        
        print(f"[DOCX DEBUG] Found {len(text_paragraphs)} paragraphs with text")
        print(f"[DOCX DEBUG] Total document paragraphs: {len(doc.paragraphs)}")
        
        if not text_paragraphs:
            print(f"[DOCX DEBUG] No extractable text found in {os.path.basename(filepath)}")
            return f"[DOCX file uploaded but no extractable text found in {os.path.basename(filepath)}]"
        
        docx_content = '\n'.join(text_paragraphs)
        print(f"[DOCX DEBUG] Extracted {len(docx_content)} characters before cleaning")
        
        # Apply word boundary restoration (same as PDF extraction)
        # Step 1: Add spaces between lowercase and uppercase (CamelCase)
        docx_content = re.sub(r'([a-z])([A-Z])', r'\1 \2', docx_content)
        
        # Step 2: Add spaces before common English word starts
        common_words = [
            'the', 'and', 'for', 'with', 'that', 'from', 'have', 'this', 'will', 'your', 'about', 'which', 'their', 'would', 'been', 'has', 'were', 'does', 'how', 'what', 'when', 'where', 'why', 'all', 'can', 'did', 'get', 'let', 'may', 'new', 'now',
            'old', 'one', 'our', 'out', 'own', 'say', 'she', 'too', 'use', 'way', 'who', 'boy', 'car', 'day', 'dog', 'go', 'good', 'hand', 'he', 'her', 'him',
            'is', 'it', 'just', 'know', 'like', 'look', 'make', 'man', 'many', 'me', 'my', 'no', 'not', 'of', 'on', 'only', 'or', 'people', 'see', 'some', 'such', 'take',
            'than', 'them', 'then', 'there', 'these', 'they', 'thing', 'think', 'through', 'to', 'tell', 'time', 'two', 'up', 'us', 'very', 'want', 'we', 'work', 'year',
            'you', 'after', 'also', 'as', 'at', 'be', 'because', 'before', 'by', 'come', 'could', 'down', 'each', 'even', 'find', 'first', 'follow', 'give', 'group', 'has', 'high', 'his', 'home', 'if', 'in', 'into', 'its',
            'just', 'keep', 'kind', 'last', 'life', 'little', 'live', 'long', 'love', 'much', 'must', 'name', 'never', 'next', 'number', 'off', 'often', 'other',
            'over', 'part', 'place', 'play', 'present', 'problem', 'provide', 'put', 'quality', 'question', 'quite', 'rather', 'read', 'result', 'right', 'run',
            'same', 'school', 'second', 'set', 'should', 'show', 'small', 'so', 'something', 'sometimes', 'state', 'still', 'study', 'system', 'take', 'tell', 'than',
            'that', 'the', 'their', 'them', 'these', 'they', 'thing', 'think', 'this', 'those', 'though', 'through', 'thus', 'to', 'today', 'together', 'too', 'took',
            'try', 'type', 'under', 'understand', 'unit', 'until', 'upon', 'use', 'used', 'usually', 'value', 'various', 'very', 'view', 'want', 'was', 'water', 'way',
            'week', 'well', 'went', 'were', 'what', 'when', 'where', 'which', 'while', 'who', 'whole', 'why', 'will', 'with', 'within', 'without', 'word', 'work',
            'world', 'would', 'write', 'written', 'year', 'yes', 'yet', 'you', 'young', 'your', 'yourself'
        ]
        
        for word in common_words:
            pattern = rf'([a-z\d])({word})(?=[a-z])'
            docx_content = re.sub(pattern, rf'\1 \2', docx_content, flags=re.IGNORECASE)
        
        # Step 3: Handle digits followed by letters
        docx_content = re.sub(r'(\d)([a-zA-Z])', r'\1 \2', docx_content)
        
        # Step 4: Handle letters followed by digits
        docx_content = re.sub(r'([a-zA-Z])(\d)', r'\1 \2', docx_content)
        
        # Step 5: Fix internal word spacing
        docx_content = re.sub(r'(\w)\s{1,2}(?=[a-z])', r'\1', docx_content)
        
        # Step 6: Fix excessive whitespace
        docx_content = re.sub(r'\n\s*\n', '\n\n', docx_content)
        docx_content = re.sub(r' {2,}', ' ', docx_content)
        docx_content = re.sub(r'\t', ' ', docx_content)
        
        # Step 7: Fix punctuation spacing
        docx_content = re.sub(r'\s+([.,!?;:])', r'\1', docx_content)
        docx_content = re.sub(r'([.!?])\s{1,}([A-Z])', r'\1 \2', docx_content)
        
        # Step 8: Additional aggressive patterns for research terms
        research_terms = [
            'research', 'model', 'network', 'algorithm', 'learning', 'neural', 'data', 'analysis',
            'method', 'approach', 'framework', 'system', 'performance', 'result', 'evaluation',
            'training', 'testing', 'validation', 'accuracy', 'metric', 'parameter', 'optimization',
            'implementation', 'architecture', 'layer', 'feature', 'extraction', 'classification',
            'prediction', 'inference', 'abstract', 'conclusion', 'introduction', 'experiment',
            'baseline', 'benchmark', 'dataset', 'corpus', 'annotation'
        ]
        
        for term in research_terms:
            pattern = rf'([a-z])({term})(?=[a-z])'
            docx_content = re.sub(pattern, rf'\1 \2', docx_content, flags=re.IGNORECASE)
        
        # Step 9: Very long lowercase word splitting (concatenation detection)
        words_to_check = docx_content.split()
        processed_words = []
        
        for word in words_to_check:
            if len(word) > 15 and word.islower() and not any(c.isdigit() for c in word):
                split_word = re.sub(r'(ing|tion|ness|ment|able|ible|ful|less|ly|er|or|ar|ian)(?=[a-z])', r'\1 ', word)
                if split_word != word:
                    word = split_word
            processed_words.append(word)
        
        docx_content = ' '.join(processed_words)
        
        # Step 10: Final cleanup
        docx_content = re.sub(r' {2,}', ' ', docx_content)
        docx_content = '\n'.join(line.strip() for line in docx_content.split('\n') if line.strip())
        
        return docx_content if docx_content else f"[No readable content in {os.path.basename(filepath)}]"
    
    except Exception as e:
        print(f"DOCX extraction error: {e}")
        return f"[Error extracting DOCX: {str(e)}]"

# Enhanced extractive summarizer as fallback
def simple_summarize(text, max_length=2000, min_length=700):
    """Enhanced extractive summarization generating comprehensive, meaningful abstracts"""
    import re
    
    # First, clean up the text more carefully - fix OCR artifacts without breaking normal words
    text = re.sub(r'([a-z])([A-Z])', r'\1 \2', text)  # Only insert space between camelCase
    text = re.sub(r'(\w)\s{2,}(?=[a-z])', r'\1', text)  # Fix broken words (internal spaces only)
    text = re.sub(r' {2,}', ' ', text)  # Multiple spaces to single
    text = re.sub(r'(\w)\n(?=[a-z])', r'\1 ', text)  # Join broken words across lines
    
    # Split by sentence endings
    sentences = re.split(r'[.!?]+', text)
    sentences = [re.sub(r' {2,}', ' ', s.strip()) for s in sentences if s.strip()]
    
    if not sentences:
        return text[:max_length] if len(text) > 0 else "No content to summarize"
    
    if len(sentences) == 1:
        return sentences[0][:max_length]
    
    # Filter out incomplete/fragmented sentences - be strict about quality
    valid_sentences = []
    for s in sentences:
        words = s.split()
        
        # Skip very short sentences
        if len(words) < 6:
            continue
        
        alpha_words = [w for w in words if any(c.isalpha() for c in w)]
        
        # Must have good alphabetic content (not metadata/numbers)
        if len(alpha_words) < len(words) * 0.65:
            continue
        
        special_ratio = sum(1 for c in s if c in '->:*[]()') / len(s)
        if special_ratio >= 0.20:
            continue
        
        valid_sentences.append(s)
    
    # If filtering removed too much, be more lenient
    if len(valid_sentences) < 4:
        valid_sentences = [s for s in sentences if len(s.split()) >= 6]
    
    if not valid_sentences:
        valid_sentences = sentences
    
    # Expanded stopwords list
    stopwords = {'the', 'a', 'an', 'and', 'or', 'but', 'in', 'on', 'at', 'to', 'for',
                 'of', 'is', 'was', 'are', 'be', 'by', 'this', 'that', 'with', 'as',
                 'from', 'into', 'up', 'about', 'which', 'who', 'it', 'its', 'they', 'them',
                 'have', 'has', 'had', 'do', 'does', 'did', 'will', 'would', 'could', 'should',
                 'may', 'might', 'can', 'must', 'shall', 'these', 'those', 'been', 'being',
                 'such', 'no', 'not', 'only', 'just', 'also', 'all', 'each', 'every', 'we', 'us',
                 'our', 'you', 'your', 'he', 'him', 'her', 'his', 'she', 'hers', 'any', 'i', 'me',
                 'than', 'then', 'same', 'other', 'through', 'during', 'before', 'after', 'above',
                 'below', 'through', 'among', 'between', 'including', 'where', 'when', 'what', 'why', 'how'}
    
    # Calculate word frequencies from valid content - focus on content words
    all_words = re.findall(r'\w+', ' '.join(valid_sentences).lower())
    word_freq = {}
    for word in all_words:
        if len(word) > 2 and word not in stopwords:
            word_freq[word] = word_freq.get(word, 0) + 1
    
    # Identify main themes (most frequent significant words) - top 20 for better coverage
    main_themes = sorted(word_freq.items(), key=lambda x: x[1], reverse=True)[:20]
    theme_words = set(word for word, _ in main_themes)
    theme_importance = dict(main_themes)
    
    # Score sentences based on multiple comprehensive factors
    sentence_scores = {}
    for idx, sentence in enumerate(valid_sentences):
        sentence_words = re.findall(r'\w+', sentence.lower())
        
        # Factor 1: Theme word presence with importance weighting
        theme_score = 0
        for word in sentence_words:
            if word in theme_words:
                theme_score += theme_importance.get(word, 0)
        theme_score = theme_score / max(len(sentence_words), 1)
        
        # Factor 2: Average word frequency (how important are the words used)
        if sentence_words:
            freq_score = sum(word_freq.get(word, 0) for word in sentence_words) / len(sentence_words)
        else:
            freq_score = 0
        
        # Factor 3: Sentence position (spread across document for diverse coverage)
        normalized_pos = idx / max(len(valid_sentences) - 1, 1)
        if normalized_pos < 0.2:  # First 20%
            position_score = 0.95
        elif normalized_pos < 0.7:  # Middle 50%
            position_score = 1.0
        else:  # Last 30%
            position_score = 0.8
        
        # Factor 4: Sentence length preference (reject too short, prefer substantial sentences)
        sent_length = len(sentence_words)
        if 8 <= sent_length <= 40:
            length_score = 1.0
        elif 5 <= sent_length < 8:
            length_score = 0.5
        elif sent_length > 40:
            length_score = 0.9
        else:
            length_score = 0.2
        
        # Factor 5: Keyword density (sentences with more unique meaningful terms)
        unique_keywords = [w for w in set(sentence_words) if word_freq.get(w, 0) > 0 and len(w) > 3]
        keyword_score = len(unique_keywords) / max(len(set(sentence_words)), 1)
        
        # Factor 6: Content density - how packed with meaningful terms
        content_words = [w for w in sentence_words if w not in stopwords and len(w) > 2]
        content_density = len(content_words) / max(len(sentence_words), 1)
        
        # Combined weighted score - emphasize theme and frequency for coherent abstracts
        total_score = (theme_score * 0.32) + (freq_score * 0.28) + (keyword_score * 0.18) + (position_score * 0.12) + (length_score * 0.05) + (content_density * 0.05)
        sentence_scores[idx] = total_score
    
    # Calculate target number of sentences for comprehensive summary
    avg_sentence_length = sum(len(s.split()) for s in valid_sentences) / len(valid_sentences) if valid_sentences else 15
    num_sentences = max(8, int(max_length / avg_sentence_length))  # Minimum 8 sentences for comprehensive abstract
    num_sentences = min(num_sentences, len(valid_sentences))
    
    # Ensure we get a diverse selection covering the document
    sorted_scores = sorted(sentence_scores.items(), key=lambda x: x[1], reverse=True)
    selected_indices = sorted([idx for idx, _ in sorted_scores[:num_sentences]])
    summary_sentences = [valid_sentences[i] for i in selected_indices]
    
    # Build summary with proper formatting (space joining for natural flow)
    summary = ' '.join(summary_sentences)
    
    # Ensure proper punctuation
    if summary and not summary.endswith('.'):
        summary += '.'
    
    # Truncate if needed while preserving sentence structure
    if len(summary) > max_length:
        summary = summary[:max_length]
        # Remove partial sentence at end
        if '.' in summary:
            summary = summary[:summary.rfind('.')+1]
    
    # Extend if too short - fill with next best sentences
    if len(summary) < min_length and len(selected_indices) < len(valid_sentences):
        remaining = sorted([idx for idx in range(len(valid_sentences)) if idx not in selected_indices],
                          key=lambda idx: sentence_scores.get(idx, 0), reverse=True)
        for idx in remaining:
            if len(summary) < min_length:
                summary += ' ' + valid_sentences[idx]
                if not summary.endswith('.'):
                    summary += '.'
            else:
                break
    
    return summary if summary else "Unable to generate summary from content"


# ==================== HEALTH CHECK ====================
@app.route('/api/health', methods=['GET'])
def health_check():
    """Health check endpoint"""
    qa = get_qa_system()
    return jsonify({
        'status': 'healthy',
        'message': 'Research Paper QA System API is running',
        'qa_system_ready': qa is not None
    }), 200


# ==================== SUMMARIZATION ====================
@app.route('/api/summarize', methods=['POST'])
def summarize():
    """
    Summarize research paper text with structured output.
    Expected JSON: {"text": "paper content", "use_structured": true, "total_length": 1000}
    """
    try:
        data = request.get_json()
        if not data or 'text' not in data:
            return jsonify({'error': 'Missing required field: text'}), 400
        
        text = data['text'].strip()
        
        # Only reject if text is empty
        if not text:
            return jsonify({'error': 'Text content cannot be empty'}), 400
        
        # Check if user wants structured or simple summary
        use_structured = data.get('use_structured', True)  # Default to structured
        total_length = data.get('total_length', 1000)  # For structured: total tokens
        max_length = data.get('max_length', 250)  # For simple: per-section tokens
        min_length = data.get('min_length', 100)
        
        print(f"\n[SUMMARIZE REQUEST] Input text length: {len(text)} chars")
        print(f"[Format] Structured: {use_structured}, Total length: {total_length}")
        
        summary = None
        
        # Try to use structured summarizer if requested
        if use_structured:
            qa = get_qa_system()
            if qa and hasattr(qa, 'structured_summarizer') and qa.structured_summarizer:
                try:
                    print("[SUMMARY] Using structured summarizer (Introduction, Methods, Results, Conclusion, Key Findings)")
                    summary = qa.structured_summarizer.summarize_structured(
                        text,
                        total_length=total_length
                    )
                except Exception as e:
                    print(f"[WARN] Structured summarizer failed: {e}, falling back to simple...")
                    import traceback
                    traceback.print_exc()
                    summary = None
        
        # Fallback to simple summarizer if structured not available or failed
        if summary is None:
            qa = get_qa_system()
            if qa and hasattr(qa, 'summarizer') and qa.summarizer:
                try:
                    print("[SUMMARY] Using simple summarizer (paragraph format)")
                    summary = qa.summarizer.summarize(
                        text,
                        max_length=max_length,
                        min_length=min_length,
                        format_as_points=False  # No bullet points, full paragraphs
                    )
                except Exception as e:
                    print(f"[WARN] QA System summarizer failed: {e}, using fallback...")
                    summary = simple_summarize(text, max_length, min_length)
            else:
                # Use simple fallback summarizer
                summary = simple_summarize(text, max_length, min_length)
        
        # Log the summary to terminal with full details
        print("\n" + "="*80)
        print("[GENERATED SUMMARY - MODEL OUTPUT]")
        print("="*80)
        print(summary)
        print("="*80)
        print(f"[Stats] Original: {len(text)} chars | Summary: {len(summary)} chars | Compression: {round((len(summary) / len(text)) * 100, 2)}%")
        print("="*80 + "\n")
        
        return jsonify({
            'success': True,
            'summary': summary,
            'original_length': len(text),
            'summary_length': len(summary),
            'compression_ratio': round((len(summary) / len(text)) * 100, 2) if len(text) > 0 else 0
        }), 200
    
    except Exception as e:
        return jsonify({'error': str(e)}), 500


# ==================== QUESTION ANSWERING ====================
@app.route('/api/answer', methods=['POST'])
def answer_question():
    """
    Answer questions about research paper content
    Expected JSON: {"question": "What is...", "context": "paper text"}
    """
    try:
        qa = get_qa_system()
        if not qa or not qa.qa_model:
            print("[API] QA System not ready")
            # Try to respond with a fallback
            return jsonify({
                'success': False,
                'error': 'QA System not initialized. Please ensure the system is properly loaded.',
                'message': 'The QA model is still loading. Please try again in a few seconds.'
            }), 503
        
        data = request.get_json()
        if not data:
            return jsonify({'error': 'No JSON data provided'}), 400
            
        question = data.get('question', '').strip()
        context = data.get('context', '').strip()
        
        # Validate inputs
        if not question:
            return jsonify({'error': 'Question cannot be empty'}), 400
        if not context:
            return jsonify({'error': 'Context cannot be empty'}), 400
        
        print(f"\n[API /answer] Processing question: '{question[:50]}...'")
        print(f"[API /answer] Context length: {len(context)} characters")
        
        # Get answer from QA model
        result = qa.qa_model.answer_question(question, context)
        
        print(f"[API /answer] Answer received: {result.get('answer', 'N/A')[:100]}")
        
        return jsonify({
            'success': True,
            'question': question,
            'answer': result.get('answer', 'No answer found'),
            'score': float(result.get('score', 0.0)),
            'confidence': 'high' if result.get('score', 0.0) > 0.5 else 'low',
            'full_result': result
        }), 200
    
    except Exception as e:
        print(f"[ERROR] /api/answer endpoint error: {type(e).__name__}: {e}")
        import traceback
        traceback.print_exc()
        return jsonify({
            'error': 'Internal server error',
            'message': str(e)
        }), 500


# ==================== SEMANTIC SEARCH ====================
@app.route('/api/search', methods=['POST'])
def search_papers():
    """
    Search papers using semantic similarity
    Expected JSON: {"query": "search term", "papers": [...], "top_k": 5}
    """
    try:
        qa = get_qa_system()
        if not qa:
            return jsonify({'error': 'QA System not initialized'}), 500
        
        data = request.get_json()
        if not data or 'query' not in data:
            return jsonify({'error': 'Missing required field: query'}), 400
        
        query = data['query']
        papers = data.get('papers', [])
        top_k = data.get('top_k', 5)
        
        results = qa.retriever.search(query, papers, top_k=top_k)
        
        return jsonify({
            'success': True,
            'query': query,
            'results': results,
            'count': len(results)
        }), 200
    
    except Exception as e:
        return jsonify({'error': str(e)}), 500


# ==================== DOCUMENT UPLOAD ====================
@app.route('/api/documents/upload', methods=['POST'])
def upload_document():
    """
    Upload and process a document (PDF, TXT, DOCX)
    Extracts text content and returns it
    """
    try:
        if 'file' not in request.files:
            return jsonify({'error': 'No file provided'}), 400
        
        file = request.files['file']
        if file.filename == '':
            return jsonify({'error': 'Empty filename'}), 400
        
        # Save file temporarily
        upload_folder = 'data/uploads'
        Path(upload_folder).mkdir(parents=True, exist_ok=True)
        filepath = os.path.join(upload_folder, file.filename)
        file.save(filepath)
        
        # Extract text based on file type
        content = None
        file_ext = os.path.splitext(file.filename)[1].lower()
        
        if file_ext == '.pdf':
            print(f"[PDF] Extracting text from PDF: {file.filename}")
            content = extract_pdf_text(filepath)
            if content:
                print(f"[OK] Extracted {len(content)} characters from PDF")
                print(f"[PDF CONTENT PREVIEW] {content[:200]}...")
        elif file_ext in ['.txt', '.text']:
            print(f"[TXT] Extracting text from text file: {file.filename}")
            content = extract_text_file(filepath)
            if content:
                print(f"[OK] Extracted {len(content)} characters from text file")
                print(f"[TXT CONTENT PREVIEW] {content[:200]}...")
        elif file_ext == '.docx':
            print(f"[DOCX] Extracting text from DOCX: {file.filename}")
            content = extract_docx_text(filepath)
            if content:
                print(f"[OK] Extracted {len(content)} characters from DOCX")
                print(f"[DOCX CONTENT PREVIEW] {content[:200]}...")
        
        # If no content extracted, return error or placeholder
        if not content:
            print(f"[ERROR] No content extracted from {file.filename}")
            content = f"[Document: {file.filename}]\n[{file_ext} file uploaded successfully. Content extraction not available.]"
        
        return jsonify({
            'success': True,
            'message': 'File uploaded and processed successfully',
            'filename': file.filename,
            'filepath': filepath,
            'content': content,
            'content_length': len(content)
        }), 200
    
    except Exception as e:
        print(f"❌ Upload error: {e}")
        return jsonify({'error': str(e)}), 500


# ==================== DOCX FILE SUMMARIZATION ====================
@app.route('/api/summarize-file', methods=['POST'])
def summarize_file():
    """
    Upload a DOCX/PDF/TXT file and summarize it directly
    Returns the extracted text summary
    """
    try:
        if 'file' not in request.files:
            return jsonify({'error': 'No file provided'}), 400
        
        file = request.files['file']
        if file.filename == '':
            return jsonify({'error': 'Empty filename'}), 400
        
        # Save file temporarily
        upload_folder = 'data/uploads'
        Path(upload_folder).mkdir(parents=True, exist_ok=True)
        filepath = os.path.join(upload_folder, file.filename)
        file.save(filepath)
        
        # Extract text based on file type
        content = None
        file_ext = os.path.splitext(file.filename)[1].lower()
        
        if file_ext == '.pdf':
            print(f"[PDF SUMMARIZE] Processing: {file.filename}")
            content = extract_pdf_text(filepath)
        elif file_ext in ['.txt', '.text']:
            print(f"[TXT SUMMARIZE] Processing: {file.filename}")
            content = extract_text_file(filepath)
        elif file_ext == '.docx':
            print(f"[DOCX SUMMARIZE] Processing: {file.filename}")
            content = extract_docx_text(filepath)
        else:
            return jsonify({'error': f'Unsupported file type: {file_ext}'}), 400
        
        # Validate extracted content
        if not content or content.startswith('['):
            print(f"[ERROR] Failed to extract text from {file.filename}")
            return jsonify({
                'error': f'Failed to extract text from {file.filename}',
                'details': content
            }), 400
        
        # Generate summary using the extracted text
        use_structured = request.form.get('use_structured', 'true').lower() == 'true'
        total_length = request.form.get('total_length', 1000, type=int)
        max_length = request.form.get('max_length', 250, type=int)
        min_length = request.form.get('min_length', 100, type=int)
        
        print(f"\n[SUMMARIZING EXTRACTED CONTENT]")
        print(f"[Content Length] {len(content)} characters")
        print(f"[Format] Structured: {use_structured}, Total length: {total_length}")
        
        # Generate summary
        summary = None
        qa = get_qa_system()
        
        # Try structured summarizer first if requested
        if use_structured:
            if qa and hasattr(qa, 'structured_summarizer') and qa.structured_summarizer:
                try:
                    print(f"[SUMMARY] Using structured summarizer...")
                    summary = qa.structured_summarizer.summarize_structured(
                        content,
                        total_length=total_length
                    )
                except Exception as e:
                    print(f"[WARN] Structured summarizer failed: {type(e).__name__}: {e}")
                    import traceback
                    traceback.print_exc()
                    summary = None
        
        # Fallback to simple summarizer
        if summary is None:
            if qa and hasattr(qa, 'summarizer') and qa.summarizer:
                try:
                    print(f"[SUMMARY] Using fallback simple summarizer...")
                    summary = qa.summarizer.summarize(
                        content,
                        max_length=max_length,
                        min_length=min_length,
                        format_as_points=False
                    )
                except Exception as e:
                    print(f"[WARN] QA System summarizer failed: {type(e).__name__}: {e}")
                    import traceback
                    traceback.print_exc()
                    summary = simple_summarize(content, max_length, min_length)
            else:
                print(f"[DEBUG] QA system not available, using fallback")
                summary = simple_summarize(content, max_length, min_length)
        
        # Log the summary
        print("\n" + "="*80)
        print("[GENERATED SUMMARY - MODEL OUTPUT]")
        print("="*80)
        print(summary)
        print("="*80)
        print(f"[Stats] Original: {len(content)} chars | Summary: {len(summary)} chars | Compression: {round((len(summary) / len(content)) * 100, 2)}%")
        print("="*80 + "\n")
        
        # Clean up uploaded file
        try:
            os.remove(filepath)
        except:
            pass
        
        return jsonify({
            'success': True,
            'filename': file.filename,
            'file_type': file_ext,
            'extracted_text': content,
            'summary': summary,
            'original_length': len(content),
            'summary_length': len(summary),
            'compression_ratio': round((len(summary) / len(content)) * 100, 2) if len(content) > 0 else 0
        }), 200
    
    except Exception as e:
        print(f"❌ File summarization error: {e}")
        return jsonify({'error': str(e)}), 500


# ==================== ERROR HANDLERS ====================
@app.errorhandler(404)
def not_found(error):
    return jsonify({'error': 'Endpoint not found'}), 404


@app.errorhandler(500)
def internal_error(error):
    return jsonify({'error': 'Internal server error'}), 500


if __name__ == '__main__':
    print("🚀 Starting Research Paper QA System API Server...")
    print("📍 Server running at: http://localhost:5000")
    print("📚 API Documentation available at: http://localhost:5000/api/health")
    app.run(debug=True, host='0.0.0.0', port=5000, use_reloader=False)
